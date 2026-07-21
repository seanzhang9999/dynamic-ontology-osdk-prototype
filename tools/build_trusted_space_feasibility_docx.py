from __future__ import annotations

import math
import re
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/research/2026-07-19-osdk-agent-trusted-data-space-feasibility.md"
OUT_DOCX = ROOT / "docs/research/2026-07-19-osdk-agent-trusted-data-space-feasibility.docx"
ASSET_DIR = ROOT / "docs/assets/research/osdk-agent-trusted-data-space-feasibility"

FONT_CANDIDATES = [
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    "/Library/Fonts/Arial Unicode.ttf",
    "/System/Library/Fonts/STHeiti Medium.ttc",
    "/System/Library/Fonts/Supplemental/Songti.ttc",
]


@dataclass
class Edge:
    source: str
    target: str


def font_path() -> str:
    for candidate in FONT_CANDIDATES:
        if Path(candidate).exists():
            return candidate
    raise FileNotFoundError("No usable Chinese font found")


FONT_PATH = font_path()


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_PATH, size=size)


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = text.replace("<br/>", "\n")
    return text.strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "Arial Unicode MS"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def set_doc_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    normal = document.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        style = document.styles[style_name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")


def wrap_text(text: str, max_chars: int) -> list[str]:
    lines: list[str] = []
    for raw in text.replace("<br/>", "\n").split("\n"):
        raw = raw.strip()
        if not raw:
            continue
        while len(raw) > max_chars:
            lines.append(raw[:max_chars])
            raw = raw[max_chars:]
        lines.append(raw)
    return lines or [""]


def draw_rounded_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    text: str,
    *,
    fill: str = "#F6FAFF",
    outline: str = "#4F7FD8",
    text_color: str = "#0F2544",
    font_size: int = 26,
) -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    font = load_font(font_size)
    lines = wrap_text(text, max(8, (x2 - x1) // 26))
    line_height = font_size + 8
    total_height = line_height * len(lines)
    y = y1 + max(8, (y2 - y1 - total_height) // 2)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        w = bbox[2] - bbox[0]
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill=text_color)
        y += line_height


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#52657A") -> None:
    draw.line([start, end], fill=color, width=3)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_len = 18
    arrow_angle = math.pi / 7
    points = [
        end,
        (
            int(end[0] - arrow_len * math.cos(angle - arrow_angle)),
            int(end[1] - arrow_len * math.sin(angle - arrow_angle)),
        ),
        (
            int(end[0] - arrow_len * math.cos(angle + arrow_angle)),
            int(end[1] - arrow_len * math.sin(angle + arrow_angle)),
        ),
    ]
    draw.polygon(points, fill=color)


def parse_flowchart(code: str) -> tuple[str, dict[str, str], list[Edge]]:
    lines = [line.rstrip() for line in code.splitlines() if line.strip()]
    direction = "LR"
    header = lines[0].strip()
    if header.startswith("flowchart"):
        parts = header.split()
        if len(parts) > 1:
            direction = parts[1]
    labels: dict[str, str] = {}
    edges: list[Edge] = []
    node_pattern = re.compile(r'([A-Za-z0-9_]+)(?:\["(.+?)"\])?')
    for line in lines[1:]:
        if "-->" not in line:
            continue
        left, right = [part.strip() for part in line.split("-->", 1)]
        left_match = node_pattern.match(left)
        right_match = node_pattern.match(right)
        if not left_match or not right_match:
            continue
        source, source_label = left_match.group(1), left_match.group(2)
        target, target_label = right_match.group(1), right_match.group(2)
        labels.setdefault(source, source_label or source)
        labels.setdefault(target, target_label or target)
        if source_label:
            labels[source] = source_label
        if target_label:
            labels[target] = target_label
        edges.append(Edge(source, target))
    return direction, labels, edges


def topological_layers(nodes: Iterable[str], edges: list[Edge]) -> dict[str, int]:
    layer = {node: 0 for node in nodes}
    for _ in range(len(layer) + 1):
        changed = False
        for edge in edges:
            next_layer = layer[edge.source] + 1
            if layer[edge.target] < next_layer:
                layer[edge.target] = next_layer
                changed = True
        if not changed:
            break
    return layer


def render_flowchart(code: str, output: Path) -> None:
    direction, labels, edges = parse_flowchart(code)
    layers = topological_layers(labels, edges)
    groups: dict[int, list[str]] = {}
    for node, layer in layers.items():
        groups.setdefault(layer, []).append(node)
    for values in groups.values():
        values.sort()

    if direction == "TB":
        node_w, node_h = 760, 104
        margin_x, margin_y, gap_y = 90, 90, 80
        width = node_w + margin_x * 2
        height = margin_y * 2 + len(labels) * node_h + (len(labels) - 1) * gap_y
        positions: dict[str, tuple[int, int, int, int]] = {}
        ordered = []
        for layer in sorted(groups):
            ordered.extend(groups[layer])
        for idx, node in enumerate(ordered):
            y = margin_y + idx * (node_h + gap_y)
            positions[node] = (margin_x, y, margin_x + node_w, y + node_h)
        box_font_size = 32
    elif direction == "LR" and (max(groups) if groups else 0) >= 4:
        node_w, node_h = 430, 120
        margin_x, margin_y, gap_x, gap_y = 80, 80, 90, 48
        columns = 4
        max_layer = max(groups) if groups else 0
        bands = max_layer // columns + 1
        max_rows = max(len(v) for v in groups.values()) if groups else 1
        band_inner_h = max_rows * node_h + (max_rows - 1) * gap_y
        band_gap = 130
        width = margin_x * 2 + columns * node_w + (columns - 1) * gap_x
        height = margin_y * 2 + bands * band_inner_h + (bands - 1) * band_gap
        positions = {}
        for layer, nodes in groups.items():
            band = layer // columns
            column = layer % columns
            band_y = margin_y + band * (band_inner_h + band_gap)
            column_height = len(nodes) * node_h + (len(nodes) - 1) * gap_y
            start_y = band_y + max(0, (band_inner_h - column_height) // 2)
            x = margin_x + column * (node_w + gap_x)
            for row, node in enumerate(nodes):
                y = start_y + row * (node_h + gap_y)
                positions[node] = (x, y, x + node_w, y + node_h)
        box_font_size = 34
    else:
        node_w, node_h = 360, 96
        margin_x, margin_y, gap_x, gap_y = 80, 80, 100, 42
        max_layer = max(groups) if groups else 0
        max_rows = max(len(v) for v in groups.values()) if groups else 1
        width = margin_x * 2 + (max_layer + 1) * node_w + max_layer * gap_x
        height = margin_y * 2 + max_rows * node_h + (max_rows - 1) * gap_y
        positions = {}
        for layer, nodes in groups.items():
            column_height = len(nodes) * node_h + (len(nodes) - 1) * gap_y
            start_y = margin_y + max(0, (height - margin_y * 2 - column_height) // 2)
            x = margin_x + layer * (node_w + gap_x)
            for row, node in enumerate(nodes):
                y = start_y + row * (node_h + gap_y)
                positions[node] = (x, y, x + node_w, y + node_h)
        box_font_size = 28

    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    for edge in edges:
        s = positions[edge.source]
        t = positions[edge.target]
        if direction == "TB":
            start = ((s[0] + s[2]) // 2, s[3])
            end = ((t[0] + t[2]) // 2, t[1])
        else:
            start = (s[2], (s[1] + s[3]) // 2)
            end = (t[0], (t[1] + t[3]) // 2)
        draw_arrow(draw, start, end)
    palette = ["#F5F9FF", "#F9FFF8", "#FFF9F1", "#FBF7FF"]
    for idx, node in enumerate(sorted(labels)):
        draw_rounded_box(
            draw,
            positions[node],
            labels[node],
            fill=palette[idx % len(palette)],
            outline="#6EA1E8",
            font_size=box_font_size,
        )
    image.save(output)


def parse_sequence(code: str) -> tuple[list[tuple[str, str]], list[tuple[str, str, str, str]]]:
    participants: list[tuple[str, str]] = []
    messages: list[tuple[str, str, str, str]] = []
    for line in code.splitlines():
        line = line.strip()
        if not line or line == "sequenceDiagram":
            continue
        participant = re.match(r'participant\s+(\w+)\s+as\s+"(.+)"', line)
        if participant:
            participants.append((participant.group(1), participant.group(2)))
            continue
        message = re.match(r"(\w+)(-+>>)(\w+):\s+\"(.+)\"", line)
        if message:
            messages.append((message.group(1), message.group(3), message.group(4), message.group(2)))
    return participants, messages


def render_sequence(code: str, output: Path) -> None:
    participants, messages = parse_sequence(code)
    count = len(participants)
    col_w = 260
    margin_x, top, bottom = 90, 80, 80
    box_w, box_h = 220, 78
    step_y = 96
    width = margin_x * 2 + max(1, count - 1) * col_w
    height = top + box_h + 50 + len(messages) * step_y + bottom
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    font = load_font(22)
    small = load_font(19)
    x_by_id: dict[str, int] = {}
    for idx, (pid, label) in enumerate(participants):
        x = margin_x + idx * col_w
        x_by_id[pid] = x
        draw_rounded_box(
            draw,
            (x - box_w // 2, top, x + box_w // 2, top + box_h),
            label,
            fill="#F6FAFF",
            outline="#4F7FD8",
            font_size=20,
        )
        draw.line((x, top + box_h, x, height - bottom // 2), fill="#C6D3E1", width=2)
    y = top + box_h + 60
    for source, target, label, arrow in messages:
        sx, tx = x_by_id[source], x_by_id[target]
        line_y = y
        color = "#43617E"
        draw_arrow(draw, (sx, line_y), (tx, line_y), color=color)
        lines = wrap_text(label, max(14, abs(tx - sx) // 18))
        text_y = line_y - 42
        for line in lines[:2]:
            bbox = draw.textbbox((0, 0), line, font=small)
            tw = bbox[2] - bbox[0]
            draw.rectangle(
                (
                    min(sx, tx) + (abs(tx - sx) - tw) / 2 - 6,
                    text_y - 2,
                    min(sx, tx) + (abs(tx - sx) + tw) / 2 + 6,
                    text_y + 24,
                ),
                fill="#FFFFFF",
            )
            draw.text(
                (min(sx, tx) + (abs(tx - sx) - tw) / 2, text_y),
                line,
                font=small,
                fill="#0F2544",
            )
            text_y += 25
        y += step_y
    image.save(output)


def render_mermaid(code: str, index: int) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    output = ASSET_DIR / f"diagram-{index:02d}.png"
    if code.strip().startswith("sequenceDiagram"):
        render_sequence(code, output)
    else:
        render_flowchart(code, output)
    return output


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(code)
    run.font.name = "Menlo"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(40, 48, 60)


def add_picture_scaled(document: Document, image_path: Path, *, max_width: float = 6.6, max_height: float = 8.6):
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    aspect = width_px / height_px
    target_w = max_width
    target_h = target_w / aspect
    if target_h > max_height:
        target_h = max_height
        target_w = target_h * aspect
    document.add_picture(str(image_path), width=Inches(target_w), height=Inches(target_h))
    paragraph = document.paragraphs[-1]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return paragraph


def resolve_image_path(raw_path: str) -> Path:
    path = Path(raw_path.strip())
    if path.is_absolute():
        return path
    source_relative = (SOURCE.parent / path).resolve()
    if source_relative.exists():
        return source_relative
    return (ROOT / path).resolve()


def add_image(document: Document, alt: str, raw_path: str) -> None:
    image_path = resolve_image_path(raw_path)
    if not image_path.exists():
        add_paragraph(document, f"[图片缺失] {alt}: {raw_path}")
        return
    image_paragraph = add_picture_scaled(document, image_path, max_width=6.7, max_height=5.0)
    image_paragraph.paragraph_format.keep_with_next = True
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(clean_inline(alt))
    set_run_font(run, size=9, bold=True)


def split_table_row(line: str) -> list[str]:
    return [clean_inline(cell) for cell in line.strip().strip("|").split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=8 if len(value) > 40 else 9, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")
    document.add_paragraph()


def add_paragraph(document: Document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, size=10)
    return paragraph


def build_docx() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    document = Document()
    set_doc_defaults(document)
    lines = markdown.splitlines()
    i = 0
    diagram_index = 1
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            code = "\n".join(code_lines)
            if lang == "mermaid":
                image_path = render_mermaid(code, diagram_index)
                caption = document.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(f"图 {diagram_index}：由 Mermaid 图块渲染")
                set_run_font(run, size=9, bold=True)
                add_picture_scaled(document, image_path)
                diagram_index += 1
            else:
                add_code_block(document, code)
            continue
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            rows = [split_table_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_table_row(lines[i]))
                i += 1
            add_table(document, rows)
            continue
        image = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if image:
            add_image(document, image.group(1), image.group(2))
            i += 1
            continue
        if stripped.startswith("**截图 "):
            document.add_page_break()
            paragraph = add_paragraph(document, stripped)
            paragraph.paragraph_format.keep_with_next = True
            i += 1
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            if level == 1:
                paragraph = document.add_heading(text, level=0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                document.add_heading(text, level=min(level - 1, 3))
            i += 1
            continue
        if stripped.startswith("- "):
            add_paragraph(document, stripped[2:], style="List Bullet")
            i += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            add_paragraph(document, numbered.group(1), style="List Number")
            i += 1
            continue
        if stripped.startswith(">"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(clean_inline(stripped.lstrip("> ")))
            set_run_font(run, size=10, bold=True)
            run.font.color.rgb = RGBColor(40, 72, 102)
            i += 1
            continue
        add_paragraph(document, stripped)
        i += 1

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT_DOCX)
    print(f"wrote {OUT_DOCX}")
    print(f"rendered {diagram_index - 1} diagrams to {ASSET_DIR}")


if __name__ == "__main__":
    build_docx()
