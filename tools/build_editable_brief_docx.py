#!/usr/bin/env python3
"""Build an editable DOCX brief from the Markdown source.

The project keeps the report content in Markdown so it can be reviewed in Git.
This script converts the supported subset of Markdown used by the brief into a
native Word document: editable headings, paragraphs, lists, tables, and code
blocks, with screenshots and diagrams embedded as real images.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BODY_FONT = "Heiti SC"
CODE_FONT = "Heiti SC"
ACCENT = RGBColor(22, 93, 255)
TEXT = RGBColor(31, 41, 55)
MUTED = RGBColor(71, 85, 105)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, font: str = BODY_FONT, size: float | None = None) -> None:
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font)
    r_fonts.set(qn("w:hint"), "eastAsia")


def set_style_font(style, font: str = BODY_FONT, size: float | None = None) -> None:
    style.font.name = font
    if size is not None:
        style.font.size = Pt(size)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font)
    r_fonts.set(qn("w:hint"), "eastAsia")


def set_document_defaults(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, BODY_FONT, 10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Heading 1", 18, ACCENT),
        ("Heading 2", 14, RGBColor(15, 23, 42)),
        ("Heading 3", 11.5, RGBColor(30, 64, 175)),
    ]:
        style = styles[name]
        set_style_font(style, BODY_FONT, size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style_font(style, BODY_FONT, 10)
        style.paragraph_format.space_after = Pt(3)

    # Make Word/LibreOffice prefer an East Asian-capable font for all default runs.
    doc_defaults = styles.element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.element.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), BODY_FONT)
    rfonts.set(qn("w:hint"), "eastAsia")
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def clean_inline(text: str) -> str:
    text = text.replace("<br/>", " / ").replace("<br>", " / ")
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "")
    return text


def add_paragraph(doc: Document, text: str, style: str | None = None, bold: bool = False) -> None:
    para = doc.add_paragraph(style=style)
    for idx, part in enumerate(clean_inline(text).split("\n")):
        if idx:
            para.add_run().add_break()
        run = para.add_run(part)
        set_run_font(run)
        run.bold = bold


def add_title(doc: Document, title: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(title)
    set_run_font(run, BODY_FONT, 24)
    run.bold = True
    run.font.color.rgb = ACCENT
    para.paragraph_format.space_after = Pt(8)


def add_code_block(doc: Document, code: str, language: str) -> None:
    if language == "mermaid":
        label = doc.add_paragraph()
        label_run = label.add_run("Mermaid 源图（Markdown 中可直接渲染）")
        set_run_font(label_run, BODY_FONT, 9)
        label_run.bold = True
        label_run.font.color.rgb = MUTED
        label.paragraph_format.space_after = Pt(2)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    set_cell_margins(cell, top=110, start=120, bottom=110, end=120)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(code.rstrip() or " ")
    set_run_font(run, CODE_FONT, 8.2)
    run.font.color.rgb = RGBColor(17, 24, 39)
    doc.add_paragraph()


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [clean_inline(cell.strip()) for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "EAF2FF")
            text = row[c_idx] if c_idx < len(row) else ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(text)
            set_run_font(run, BODY_FONT, 8.8 if col_count >= 4 else 9.2)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(30, 64, 175)
    doc.add_paragraph()


def add_image(doc: Document, md_path: Path, alt: str, source_path: str) -> None:
    img_path = (md_path.parent / source_path).resolve()
    if not img_path.exists():
        add_paragraph(doc, f"[图片缺失] {alt}: {source_path}")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width = Inches(6.25)
    if "architecture-diagrams" in source_path:
        width = Inches(5.55)
    run = para.add_run()
    run.add_picture(str(img_path), width=width)
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption.add_run(f"图：{alt}")
    set_run_font(cap_run, BODY_FONT, 8.5)
    cap_run.font.color.rgb = MUTED
    caption.paragraph_format.space_after = Pt(10)


def is_image_line(line: str) -> re.Match[str] | None:
    return re.match(r"^\s*!\[([^\]]*)\]\(([^)]+)\)\s*$", line)


def is_unordered_list(line: str) -> bool:
    return bool(re.match(r"^\s*-\s+", line))


def is_ordered_list(line: str) -> bool:
    return bool(re.match(r"^\s*\d+\.\s+", line))


def is_special_start(line: str, next_line: str | None = None) -> bool:
    if not line.strip():
        return True
    if line.startswith("```"):
        return True
    if re.match(r"^#{1,6}\s+", line):
        return True
    if is_image_line(line):
        return True
    if is_unordered_list(line) or is_ordered_list(line):
        return True
    if line.strip().startswith(">"):
        return True
    if "|" in line and next_line is not None and is_table_separator(next_line):
        return True
    return False


def render_markdown(md_path: Path, out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    set_document_defaults(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    first_heading = True

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        next_line = lines[i + 1] if i + 1 < len(lines) else None

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            language = stripped.strip("`").strip()
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, "\n".join(block), language)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            text = clean_inline(heading.group(2))
            if level == 1 and first_heading:
                add_title(doc, text)
                first_heading = False
            else:
                para = doc.add_heading(text, level=min(level, 3))
                for run in para.runs:
                    set_run_font(run, BODY_FONT)
            i += 1
            continue

        img = is_image_line(line)
        if img:
            add_image(doc, md_path, img.group(1), img.group(2))
            i += 1
            continue

        if "|" in line and next_line is not None and is_table_separator(next_line):
            table_lines = [line]
            i += 2
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            rows = [split_table_row(row) for row in table_lines]
            add_table(doc, rows)
            continue

        if is_unordered_list(line) or is_ordered_list(line):
            ordered = is_ordered_list(line)
            style = "List Number" if ordered else "List Bullet"
            while i < len(lines):
                current = lines[i]
                if ordered and not is_ordered_list(current):
                    break
                if not ordered and not is_unordered_list(current):
                    break
                text = re.sub(r"^\s*(?:-|\d+\.)\s+", "", current)
                add_paragraph(doc, text, style=style)
                i += 1
            continue

        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().lstrip(">").strip())
                i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.18)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(9)
            run = para.add_run(" ".join(quote_lines))
            set_run_font(run, BODY_FONT, 10.3)
            run.italic = True
            run.font.color.rgb = RGBColor(30, 64, 175)
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            nxt_next = lines[i + 1] if i + 1 < len(lines) else None
            if is_special_start(nxt, nxt_next):
                break
            paragraph_lines.append(nxt)
            i += 1
        text = " ".join(chunk.strip().rstrip("  ") for chunk in paragraph_lines)
        add_paragraph(doc, text)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--source",
        type=Path,
        default=Path("docs/dynamic-ontology-osdk-demo-architecture-brief.md"),
    )
    parser.add_argument(
        "--out",
        type=Path,
        default=Path("docs/dynamic-ontology-osdk-demo-architecture-brief.docx"),
    )
    args = parser.parse_args()
    render_markdown(args.source.resolve(), args.out.resolve())
    print(args.out.resolve())


if __name__ == "__main__":
    main()
